"""consultant_profiles_v2.docx を解析して ConsultantProfile のリスト/JSON を生成する。

docx のレイアウトは固定でないため、ラベル付き行（例:「前職：リクルート」）と
2列テーブル（ラベル|値）を最大限拾う、寛容なパーサとして実装。
解析後は config/consultants.json に保存し、必ず内容を目視確認すること。
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from .config import scout_rules
from .models import ConsultantProfile

_LABEL_MAP: list[tuple[str, str]] = [
    (r"(?:表示名|氏名|名前|イニシャル)", "display_name"),
    (r"(?:前職企業|前職|職歴|経歴企業|出身企業)", "former_companies"),
    (r"(?:業界|業種)", "industries"),
    (r"(?:出身大学|大学|学歴)", "universities"),
    (r"(?:職種|役割|ロール)", "roles"),
    (r"(?:得意領域|専門領域|専門|スキル)", "specialties"),
    (r"(?:紹介URL|プロフィールURL|URL|ページ)", "profile_url"),
    (r"(?:タグ|分類)", "tags"),
    (r"(?:ID|社員番号)", "id"),
]

_LIST_FIELDS = {"former_companies", "industries", "universities", "roles", "specialties", "tags"}


def _split(value: str) -> list[str]:
    return [p.strip() for p in re.split(r"[、,;／/・\n]+", value) if p.strip()]


def _match_label(text: str) -> tuple[str, str] | None:
    m = re.match(r"\s*(.+?)\s*[:：]\s*(.*)$", text)
    if not m:
        return None
    label, value = m.group(1), m.group(2)
    for label_re, field in _LABEL_MAP:
        if re.fullmatch(label_re, label) or re.search(label_re, label):
            return field, value
    return None


def _infer_tags(profile: dict, rules: dict) -> list[str]:
    tags = set(profile.get("tags", []))
    cfg = rules.get("matching", {})
    blob = " ".join(profile.get("former_companies", []) + profile.get("industries", []))
    if any(k in blob for k in cfg.get("recruit_keywords", [])):
        tags.add("recruit")
    if any(k in blob for k in cfg.get("insurance_keywords", [])):
        tags.add("insurance")
    return sorted(tags)


def _finalize(profile: dict, idx: int, rules: dict) -> ConsultantProfile | None:
    if not any(profile.get(f) for f in ("display_name", "former_companies", "profile_url")):
        return None
    profile.setdefault("id", f"c{idx:03d}")
    profile.setdefault("display_name", profile["id"])
    profile["tags"] = _infer_tags(profile, rules)
    for f in _LIST_FIELDS:
        profile.setdefault(f, [])
    return ConsultantProfile(**{k: profile.get(k) for k in ConsultantProfile.model_fields})


def parse_docx(path: str | Path) -> list[ConsultantProfile]:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    rules = scout_rules()
    consultants: list[ConsultantProfile] = []
    current: dict = {}
    idx = 1

    def flush() -> None:
        nonlocal current, idx
        prof = _finalize(current, idx, rules)
        if prof:
            consultants.append(prof)
            idx += 1
        current = {}

    def apply(field: str, value: str) -> None:
        # display_name が再出現したら新しいコンサルタントの開始とみなす。
        if field == "display_name" and current.get("display_name"):
            flush()
        if field in _LIST_FIELDS:
            current.setdefault(field, []).extend(_split(value))
        else:
            current[field] = value.strip()

    # 段落
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        kv = _match_label(text)
        if kv:
            apply(*kv)
        elif para.style and para.style.name and para.style.name.startswith("Heading"):
            # 見出しは新コンサルタントの表示名として扱う。
            if current.get("display_name"):
                flush()
            current["display_name"] = text

    # テーブル（2列をラベル/値として解釈）
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                kv = _match_label(f"{cells[0]}：{cells[1]}")
                if kv:
                    apply(*kv)

    flush()
    return consultants


def import_to_json(docx_path: str | Path, out_path: str | Path) -> int:
    consultants = parse_docx(docx_path)
    payload = {
        "_note": f"{Path(docx_path).name} から自動生成。内容を必ず確認してください。",
        "consultants": [c.model_dump() for c in consultants],
    }
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    Path(out_path).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return len(consultants)
